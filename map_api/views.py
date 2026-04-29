from django.http import JsonResponse, FileResponse
from django.views.decorators.csrf import csrf_exempt
from django.conf import settings
from django.shortcuts import render
from http import HTTPStatus
import dashscope
import json
import os
import uuid
import math
import threading
import requests
from PIL import Image

from .utils.get_satellite_image import fetch_satellite_image, haversine_distance, get_download_progress, _download_progress
from .utils.image_preprocessor import smart_prepare_image, MAX_DIM_MAP
from .models import ChatHistory

# 规范化保存目录
SAVE_DIR = os.path.join(settings.MEDIA_ROOT, 'satellite_imgs')
if not os.path.exists(SAVE_DIR):
    os.makedirs(SAVE_DIR)

def index_view(request):
    """负责展示前端地图页面"""
    return render(request, 'browser.html')

# ----------------------
# 卫星图下载接口
# ----------------------
@csrf_exempt
def get_satellite_img_api(request):
    if request.method != 'POST':
        return JsonResponse({'code':405,'msg':'仅支持POST','data':None},status=405)

    try:
        data = json.loads(request.body)
        min_lng = float(data['min_lng'])
        min_lat = float(data['min_lat'])
        max_lng = float(data['max_lng'])
        max_lat = float(data['max_lat'])
        resolution = int(data.get('target_resolution', 0))

        lon_span = max_lng - min_lng
        lat_span = max_lat - min_lat
        center_lat_rad = math.radians((min_lat + max_lat) / 2)
        diag_m = haversine_distance(min_lng, min_lat, max_lng, max_lat)
        if resolution <= 0:
            if diag_m > 20000:
                resolution = 3072
            elif diag_m > 5000:
                resolution = 2048
            else:
                resolution = 1280

        # Pre-calculate tile count for progress
        MAX_TOTAL = 4096
        CELL_MAX = 1280
        target = min(MAX_TOTAL, max(1, resolution))
        aspect_ratio = (lon_span * math.cos(center_lat_rad)) / lat_span if lat_span else 1
        if aspect_ratio >= 1:
            total_w = target
            total_h = max(1, int(target / aspect_ratio))
        else:
            total_h = target
            total_w = max(1, int(target * aspect_ratio))
        if total_w <= CELL_MAX and total_h <= CELL_MAX:
            total_tiles = 1
        else:
            cols = math.ceil(total_w / CELL_MAX)
            rows = math.ceil(total_h / CELL_MAX)
            total_tiles = cols * rows

        file_name = f"sat_{uuid.uuid4().hex[:8]}.jpg"
        gsd_lon = (lon_span / resolution) * 111320 * math.cos(center_lat_rad)
        gsd_lat_val = (lat_span / (resolution / ((lon_span * math.cos(center_lat_rad)) / lat_span if lat_span else 1))) * 110574
        gsd = (gsd_lon + gsd_lat_val) / 2
        area_km2 = (lon_span * 111320 * math.cos(center_lat_rad)) * (lat_span * 110574) / 1e6

        _download_progress[file_name] = {"total": total_tiles, "done": 0, "status": "downloading"}

        def _download():
            try:
                img_path = fetch_satellite_image(
                    min_lng, min_lat, max_lng, max_lat,
                    save_dir=SAVE_DIR,
                    file_name=file_name,
                    target_resolution=resolution
                )
                if not img_path:
                    _download_progress[file_name]["status"] = "error"
            except Exception:
                _download_progress[file_name]["status"] = "error"

        threading.Thread(target=_download, daemon=True).start()

        return JsonResponse({
            "code": 200,
            "msg": "下载已启动",
            "data": {
                "file_name": file_name,
                "total_tiles": total_tiles,
                "gsd_m": round(gsd, 2),
                "area_km2": round(area_km2, 4),
                "resolution_px": resolution,
            }
        })

    except Exception as e:
        return JsonResponse({"code":400,"msg":str(e),"data":None},status=400)

# ----------------------
# 精准读取图片接口 (防缓存、防串联)
# ----------------------
def show_satellite_image(request):
    file_name = request.GET.get('file')
    if file_name:
        target_path = os.path.join(SAVE_DIR, file_name)
        if not os.path.exists(target_path):
             return JsonResponse({"code": 404, "msg": "找不到指定的卫星图"}, status=404)
    else:
        # 兼容旧版的后备逻辑
        files = sorted([f for f in os.listdir(SAVE_DIR) if f.endswith('.jpg')],
                        key=lambda x: os.path.getmtime(os.path.join(SAVE_DIR, x)), reverse=True)
        if not files:
            return JsonResponse({"code":404,"msg":"无图片"},status=404)
        target_path = os.path.join(SAVE_DIR, files[0])
        
    return FileResponse(open(target_path,'rb'), content_type='image/jpeg')

# ----------------------
# AI 多轮视觉推理接口 (Qwen3.5-Plus)
# ----------------------
@csrf_exempt
def ai_query_region(request):
    try:
        data = json.loads(request.body)
        file_name = data.get("file_name")
        file_names = data.get("file_names")
        question = data.get("question")
        front_history = data.get("history", [])
        model_name = data.get("model", "qwen3-vl-plus")

        dashscope.api_key = os.environ.get('DASHSCOPE_API_KEY', '')

        SYSTEM_PROMPT = """你是SatelliteSense，顶级的遥感图像分析专家。精通地理学、城市规划、农学、水文学等多领域。

## 回答原则
- 专业客观，使用遥感标准术语
- 用数据说话（如"约30%植被覆盖"）
- 根据问题灵活组织：开放性分析可分段阐述，具体问题直接精准回答
- 中文回答，简洁有力"""

        max_dim = MAX_DIM_MAP.get(model_name, 2560)

        # 多图对比
        if file_names and isinstance(file_names, list) and len(file_names) >= 2:
            labels = "ABCDEFGH"
            content_parts = []
            for i, fn in enumerate(file_names):
                fp = os.path.join(SAVE_DIR, fn)
                if os.path.exists(fp):
                    label = labels[i] if i < len(labels) else f"区域{i+1}"
                    content_parts.append({"image": f"file://{fp}"})
                    content_parts.append({"text": f"这是区域 {label}。"})
            if len([p for p in content_parts if "image" in p.values()]) < 2:
                return JsonResponse({"code": 400, "msg": "至少需要 2 张有效图片", "data": None})
            content_parts.append({"text": f"{SYSTEM_PROMPT}\n\n共有 {len(file_names)} 个区域的遥感卫星影像，请对比分析。"})
            messages = [{"role": "user", "content": content_parts}]
        # 单图
        else:
            if not file_name:
                return JsonResponse({"code": 400, "msg": "缺少图片标识", "data": None})
            target_path = os.path.join(SAVE_DIR, file_name)
            if not os.path.exists(target_path):
                return JsonResponse({"code": 404, "msg": "卫星图文件已丢失，请重新框选", "data": None})

            spatial_ctx = data.get("spatial_context", "")
            preprocess = smart_prepare_image(target_path, max_dim=max_dim)
            if not preprocess:
                return JsonResponse({"code": 500, "msg": "图像预处理失败", "data": None})

            if "single" in preprocess:
                effective_spatial = ""
                if preprocess["eff_w"] != preprocess["orig_w"]:
                    ratio_px = preprocess["eff_w"] / preprocess["orig_w"]
                    effective_spatial = f"（原始 {preprocess['orig_w']}×{preprocess['orig_h']} px，已优化缩放至 {preprocess['eff_w']}×{preprocess['eff_h']} px）"
                first_msg = SYSTEM_PROMPT + "\n\n---\n\n请分析这张遥感卫星影像。"
                if spatial_ctx:
                    first_msg += "\n\n## 空间上下文（供参考）\n" + spatial_ctx
                if effective_spatial:
                    first_msg += "\n" + effective_spatial
                messages = [{
                    "role": "user",
                    "content": [{"image": f"file://{preprocess['single']}"}, {"text": first_msg}]
                }]
            else:
                tiles = preprocess["tiles"]
                grid = preprocess["grid"]
                content_parts = []
                overview = tiles[0]
                content_parts.append({"image": f"file://{overview}"})
                content_parts.append({"text": "这是该区域的概览缩略图。"})
                labels = "ABCDEFGHIJKLMNOP"
                for i, tp in enumerate(tiles[1:]):
                    label = labels[i] if i < len(labels) else str(i + 1)
                    content_parts.append({"image": f"file://{tp}"})
                    content_parts.append({"text": f"这是分块 {label}（{grid[0]}×{grid[1]} 网格中的一块，{preprocess['eff_w']}×{preprocess['eff_h']} px 原始分辨率）。"})
                first_msg = SYSTEM_PROMPT
                first_msg += f"\n\n---\n\n这是一张超大遥感影像（原始 {preprocess['orig_w']}×{preprocess['orig_h']} px），已分割为 1 张概览图 + {len(tiles)-1} 个 {preprocess['eff_w']}×{preprocess['eff_h']} px 分块（{grid[0]}×{grid[1]} 网格）。请结合概览图的整体布局和各分块的原始细节，综合分析这片区域。"
                if spatial_ctx:
                    first_msg += "\n\n## 空间上下文（供参考）\n" + spatial_ctx
                content_parts.append({"text": first_msg})
                messages = [{"role": "user", "content": content_parts}]

        for msg in front_history:
            role = "user" if msg["role"] == "user" else "assistant"
            messages.append({"role": role, "content": [{"text": msg["content"]}]})

        messages.append({
            "role": "user",
            "content": [{"text": question}]
        })

        call_kwargs = {"model": model_name, "messages": messages}
        if model_name in ('qwen3-vl-plus', 'qwen3-vl-flash', 'qwen-vl-max', 'qwen-vl-plus'):
            call_kwargs["vl_high_resolution_images"] = True
        response = dashscope.MultiModalConversation.call(**call_kwargs)

        if response.status_code == HTTPStatus.OK:
            content = response.output.choices[0].message.content
            if isinstance(content, list):
                real_answer = content[0].get('text', '')
            else:
                real_answer = content

            return JsonResponse({
                "code": 200,
                "msg": "success",
                "data": {"answer": real_answer}
            })
        else:
            return JsonResponse({"code": 500, "msg": f"AI 调用失败: {response.message}", "data": None}, status=500)

    except Exception as e:
        return JsonResponse({"code": 400, "msg": f"系统异常: {str(e)}", "data": None}, status=400)


# ----------------------
# 地理搜索（高德 POI 搜索，国内网络原生支持）
# ----------------------
AMAP_KEY = "7cac8b19c3d2009ed9042836f2ca4ead"

def geo_search(request):
    q = request.GET.get('q', '').strip()
    if len(q) < 1:
        return JsonResponse({"code": 400, "msg": "请输入搜索关键词", "data": []})

    if not AMAP_KEY:
        return JsonResponse({"code": 500, "msg": "请先在 views.py 中配置 AMAP_KEY（免费获取: https://lbs.amap.com/）", "data": []})

    try:
        proxies = {"http": None, "https": None}
        url = f"https://restapi.amap.com/v3/place/text?keywords={q}&key={AMAP_KEY}&offset=8&extensions=base"
        resp = requests.get(url, timeout=8, proxies=proxies)
        data = resp.json()
        results = []
        for poi in data.get("pois", []):
            loc = poi.get("location", "0,0").split(",")
            results.append({
                "name": poi.get("name", ""),
                "display_name": f"{poi.get('pname', '')}{poi.get('cityname', '')}{poi.get('adname', '')}{poi.get('address', '')}",
                "lat": float(loc[1]) if len(loc) == 2 else 0,
                "lon": float(loc[0]) if len(loc) == 2 else 0,
                "type": poi.get("typecode", ""),
            })
        return JsonResponse({"code": 200, "data": results})
    except Exception as e:
        return JsonResponse({"code": 500, "msg": str(e), "data": []})


# ----------------------
# 下载进度查询
# ----------------------
def get_progress(request):
    file_name = request.GET.get("file", "")
    info = get_download_progress(file_name)
    if not info:
        return JsonResponse({"code": 404, "msg": "no progress found"})
    return JsonResponse({"code": 200, "data": info})


# ----------------------
# 聊天历史 CRUD
# ----------------------
@csrf_exempt
def chat_history_list(request):
    if request.method == "GET":
        histories = ChatHistory.objects.values(
            "id", "image_file", "spatial_context", "bbox", "created_at", "updated_at"
        )[:20]
        return JsonResponse({"code": 200, "data": list(histories)})
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            image_file = data.get("image_file", "")
            obj, created = ChatHistory.objects.update_or_create(
                image_file=image_file,
                defaults={
                    "messages": data.get("messages", []),
                    "spatial_context": data.get("spatial_context", ""),
                    "bbox": data.get("bbox", None),
                }
            )
            return JsonResponse({"code": 200, "data": {"id": obj.id, "created": created}})
        except Exception as e:
            return JsonResponse({"code": 400, "msg": str(e)})
    return JsonResponse({"code": 405, "msg": "method not allowed"}, status=405)


def chat_history_detail(request, history_id):
    if request.method == "GET":
        try:
            obj = ChatHistory.objects.get(id=history_id)
            return JsonResponse({"code": 200, "data": {
                "id": obj.id, "image_file": obj.image_file,
                "messages": obj.messages, "spatial_context": obj.spatial_context,
                "bbox": obj.bbox, "created_at": obj.created_at.isoformat(),
            }})
        except ChatHistory.DoesNotExist:
            return JsonResponse({"code": 404, "msg": "not found"}, status=404)
    if request.method == "DELETE":
        ChatHistory.objects.filter(id=history_id).delete()
        return JsonResponse({"code": 200, "msg": "deleted"})
    return JsonResponse({"code": 405, "msg": "method not allowed"}, status=405)


# ----------------------
# 分析报告生成（Word）
# ----------------------
@csrf_exempt
def generate_report(request):
    try:
        data = json.loads(request.body)
        file_name = data.get("file_name", "")
        title = data.get("title", "遥感分析报告")
        messages = data.get("messages", [])
        spatial_ctx = data.get("spatial_context", "")
        bbox = data.get("bbox", {})

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import datetime

        doc = Document()

        for section in doc.sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        style = doc.styles['Normal']
        style.font.name = 'Microsoft YaHei'
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.25

        for i in range(1, 4):
            hs = doc.styles[f'Heading {i}']
            hs.font.name = 'Microsoft YaHei'
            hs.font.color.rgb = RGBColor(30, 30, 30)
            hs.font.bold = False
            if i == 1:
                hs.font.size = Pt(22)
            elif i == 2:
                hs.font.size = Pt(14)
            else:
                hs.font.size = Pt(12)

        t = doc.styles['Title']
        t.font.name = 'Microsoft YaHei'
        t.font.size = Pt(26)
        t.font.bold = False
        t.font.color.rgb = RGBColor(30, 30, 30)

        h = doc.add_heading(title, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph(f"生成时间：{datetime.datetime.now().strftime('%Y-%m-%d %H:%M')}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if bbox and bbox.get("min_lng"):
            doc.add_heading("区域坐标 / 空间数据", level=2)
            loc_text = f"经度：{bbox.get('min_lng', '?')} ~ {bbox.get('max_lng', '?')}  纬度：{bbox.get('min_lat', '?')} ~ {bbox.get('max_lat', '?')}"
            if spatial_ctx:
                loc_text += f"\n{spatial_ctx}"
            doc.add_paragraph(loc_text)
        elif spatial_ctx:
            doc.add_heading("空间数据", level=2)
            doc.add_paragraph(spatial_ctx)

        if file_name:
            img_path = os.path.join(SAVE_DIR, file_name)
            if os.path.exists(img_path):
                doc.add_heading("卫星影像", level=2)
                img = Image.open(img_path)
                max_w = Cm(14)
                doc.add_picture(img_path, width=max_w)
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if messages:
            doc.add_heading("AI 分析对话", level=2)
            for msg in messages:
                role = msg.get("role", "user")
                role_label = "用户" if role == "user" else "AI 助手"
                color = RGBColor(0, 122, 255) if role == "user" else RGBColor(60, 180, 75)
                p = doc.add_paragraph()
                run_label = p.add_run(f"[{role_label}]  ")
                run_label.font.name = 'Microsoft YaHei'
                run_label.font.size = Pt(10.5)
                run_label.font.bold = True
                run_label.font.color.rgb = color
                run_content = p.add_run(msg.get("content", ""))
                run_content.font.name = 'Microsoft YaHei'
                run_content.font.size = Pt(10.5)

        report_name = f"report_{uuid.uuid4().hex[:8]}.docx"
        report_path = os.path.join(settings.MEDIA_ROOT, report_name)
        os.makedirs(settings.MEDIA_ROOT, exist_ok=True)
        doc.save(report_path)

        download_url = f"/api/report/download/?file={report_name}"
        return JsonResponse({"code": 200, "data": {"file_name": report_name, "download_url": download_url}})
    except Exception as e:
        return JsonResponse({"code": 500, "msg": str(e)})


def download_report(request):
    file_name = request.GET.get("file", "")
    path = os.path.join(settings.MEDIA_ROOT, file_name)
    if not os.path.exists(path):
        return JsonResponse({"code": 404, "msg": "not found"}, status=404)
    return FileResponse(open(path, "rb"), as_attachment=True, filename=file_name,
                        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")